from docx import Document
import os
import sys

# Set stdout to utf-8
sys.stdout.reconfigure(encoding='utf-8')

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_FILE = os.path.join(BASE_DIR, "server", "python_modules", "TEMPLATE_MEMORIAL.docx")

if not os.path.exists(TEMPLATE_FILE):
    TEMPLATE_FILE = "c:/Users/fulvi/OneDrive/Agentes - Antigravity/Fóton/temp_analysis/server/python_modules/TEMPLATE_MEMORIAL.docx"

doc = Document(TEMPLATE_FILE)

with open("template_content.txt", "w", encoding="utf-8") as f:
    f.write("--- PARAGRAPHS ---\n")
    for p in doc.paragraphs:
        if p.text.strip():
            f.write(p.text.strip() + "\n")
            
    f.write("\n--- TABLES ---\n")
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                f.write(" | ".join(row_data) + "\n")

print("Exportado para template_content.txt")
